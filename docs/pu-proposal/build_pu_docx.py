"""Builds the Kinyo project proposal in the Purbanchal University /
Himalayan Whitehouse International College format used by the sample document.

Formatting taken from the sample:
  * Times New Roman 12 pt, 1.5 line spacing, justified body text
  * A4 portrait; 1 in top/right/bottom margins, 1.25 in left margin (binding)
  * Heading 1 centred bold 16 pt on a new page; Heading 2 and 3 left bold 14 pt
  * Captions below figures and tables, centred italic
  * Page numbers centred in the footer; front matter upper-case roman, body
    restarting at 1
  * IEEE numbered references

Run:  python build_pu_docx.py
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import content_pu as C

HERE = Path(__file__).parent
FIG_DIR = HERE / "figures"
OUT = Path.home() / "Downloads" / "Kinyo_Project_Proposal.docx"

FONT = "Times New Roman"
SIZE = Pt(12)
H1_SIZE = Pt(16)
H2_SIZE = Pt(14)
CAP_SIZE = Pt(9)      # matches the sample's Caption style


# --------------------------------------------------------------- low level --
def _set_run(run, *, bold=False, italic=False, size=SIZE, caps=False,
             color=(0, 0, 0)):
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.font.all_caps = caps
    run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)


def _fmt(p, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing=1.5, space_after=0,
         left_indent=None, hanging=None, keep_next=False, page_break=False):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.line_spacing = spacing
    if left_indent is not None:
        pf.left_indent = left_indent
    if hanging is not None:
        pf.left_indent = hanging
        pf.first_line_indent = -hanging
    pf.keep_with_next = keep_next
    pf.page_break_before = page_break
    p.alignment = align
    return p


def para(doc, text, **kw):
    p = doc.add_paragraph()
    style_kw = {k: kw.pop(k) for k in list(kw)
                if k in ("align", "spacing", "space_after", "left_indent",
                         "hanging", "keep_next", "page_break")}
    _fmt(p, **style_kw)
    _set_run(p.add_run(text), **kw)
    return p


def h1(doc, text, *, page_break=True):
    p = doc.add_paragraph(style="Heading 1")
    _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, keep_next=True,
         page_break=page_break)
    _set_run(p.add_run(text), bold=True, size=H1_SIZE)
    return p


def h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    _fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0, keep_next=True)
    _set_run(p.add_run(text), bold=True, size=H2_SIZE)
    return p


def h3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    _fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0, keep_next=True)
    _set_run(p.add_run(text), bold=True, size=H2_SIZE)
    return p


def sub(doc, text):
    """A bold run-in label used for the numbered groups inside a section."""
    p = doc.add_paragraph()
    _fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0, keep_next=True)
    _set_run(p.add_run(text), bold=True)
    return p


def bullet(doc, text, *, level=0, lead=None):
    p = doc.add_paragraph()
    _fmt(p, left_indent=Inches(0.5 + 0.35 * level), hanging=None)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.left_indent = Inches(0.5 + 0.35 * level)
    _set_run(p.add_run("•  "))
    if lead:
        _set_run(p.add_run(lead + ": "), bold=True)
    _set_run(p.add_run(text))
    return p


def numbered(doc, idx, text):
    p = doc.add_paragraph()
    _fmt(p, hanging=Inches(0.4))
    _set_run(p.add_run(f"{idx}.  {text}"))
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0, space_after=10)
    _set_run(p.add_run(text), italic=True, size=CAP_SIZE, color=(0x1F, 0x33, 0x51))
    return p


# ------------------------------------------------------------------ fields --
def _field(paragraph, instr, placeholder="1"):
    r = paragraph.add_run()
    f = OxmlElement("w:fldChar")
    f.set(qn("w:fldCharType"), "begin")
    r._element.append(f)
    r = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    r._element.append(it)
    r = paragraph.add_run()
    f = OxmlElement("w:fldChar")
    f.set(qn("w:fldCharType"), "separate")
    r._element.append(f)
    r = paragraph.add_run(placeholder)
    _set_run(r)
    r = paragraph.add_run()
    f = OxmlElement("w:fldChar")
    f.set(qn("w:fldCharType"), "end")
    r._element.append(f)


def footer_page_number(section, *, show=True):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if show:
        _field(p, " PAGE ")


def set_page_numbering(section, fmt, start=None):
    sectPr = section._sectPr
    for existing in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(existing)
    el = OxmlElement("w:pgNumType")
    el.set(qn("w:fmt"), fmt)
    if start is not None:
        el.set(qn("w:start"), str(start))
    sectPr.append(el)


# ------------------------------------------------------------------ tables --
def _shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def table(doc, headers, rows, widths, *, number, title, size=Pt(11)):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    tblPr = t._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for col, w in zip(t.columns, widths):
        col.width = Inches(w)

    def fill(cells, values, bold=False):
        for cell, value, w in zip(cells, values, widths):
            cell.width = Inches(w)
            cp = cell.paragraphs[0]
            _fmt(cp, align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, space_after=2)
            _set_run(cp.add_run(str(value)), bold=bold, size=size)

    fill(t.rows[0].cells, headers, bold=True)
    for cell in t.rows[0].cells:
        _shade(cell, "DCE4F0")
    trPr = t.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)

    for row in rows:
        fill(t.add_row().cells, row)

    if number != "":
        caption(doc, f"Table {number} {title}")
    else:
        para(doc, "", space_after=8)
    return t


def figure(doc, number):
    title, filename, width_in = C.FIGURES[number]
    p = doc.add_paragraph()
    _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0, space_after=0,
         keep_next=True)
    p.add_run().add_picture(str(FIG_DIR / filename), width=Inches(width_in))
    caption(doc, f"Figure {number} {title}")


# ------------------------------------------------------------------- build --
def build():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = SIZE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rpr = normal.element.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), FONT)
    rpr.append(rf)

    for name, size in (("Heading 1", H1_SIZE), ("Heading 2", H2_SIZE),
                       ("Heading 3", H2_SIZE)):
        st = doc.styles[name]
        st.font.name = FONT
        st.font.size = size
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)
        st.paragraph_format.line_spacing = 1.5
        if name == "Heading 1":
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        srpr = st.element.get_or_add_rPr()
        srf = OxmlElement("w:rFonts")
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            srf.set(qn(a), FONT)
        srpr.append(srf)

    sec = doc.sections[0]
    sec.page_width = Inches(8.27)
    sec.page_height = Inches(11.69)
    sec.top_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.25)
    footer_page_number(sec, show=False)

    m = C.META

    # --------------------------------------------------------- title page --
    def centre(text, *, bold=False, size=SIZE, space=0, italic=False):
        p = doc.add_paragraph()
        _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.15, space_after=space)
        _set_run(p.add_run(text), bold=bold, size=size, italic=italic)

    centre("", space=6)
    centre(m["university"], bold=True, size=Pt(16), space=2)
    centre(m["college"], bold=True, size=Pt(14), space=2)
    centre(m["college_address"], bold=True, size=Pt(12), space=24)
    centre("A Project Proposal", size=Pt(14), space=2)
    centre("On", size=Pt(12), space=10)
    centre(m["title"], bold=True, size=Pt(14), space=6)
    centre(m["semester_line"], italic=True, size=Pt(12), space=28)
    centre("Submitted By:", bold=True, space=6)
    for name, roll in m["students"]:
        centre(f"{name} [{roll}]", space=2)
    centre("", space=20)
    centre("Submitted To:", bold=True, space=6)
    centre("THE DEPARTMENT OF SCIENCE AND TECHNOLOGY", space=2)
    centre("IN PARTIAL FULFILLMENT OF THE REQUIREMENTS FOR THE DEGREE OF "
           "BACHELOR OF INFORMATION TECHNOLOGY", space=30)
    centre(m["date"], space=2)
    centre(m["city"])

    # -------------------------------------------------------- front matter --
    fm = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_numbering(fm, "upperRoman", start=1)
    footer_page_number(fm, show=True)

    h1(doc, "ABSTRACT", page_break=False)
    for t in C.ABSTRACT[:-1]:
        para(doc, t)
    para(doc, C.ABSTRACT[-1], space_after=0)

    h1(doc, "TABLE OF CONTENTS")
    p = doc.add_paragraph()
    _fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT)
    _field(p, r' TOC \o "1-3" \h \z \u ', "Right-click and choose Update Field.")

    h1(doc, "LIST OF FIGURES")
    table(doc,
          ["Figure", "Title", "Page"],
          [[f"Figure {n}", C.FIGURES[n][0], ""] for n in sorted(C.FIGURES)],
          [1.0, 3.5, 1.0], number="", title="",
          )
    h1(doc, "LIST OF TABLES")
    table(doc,
          ["Table", "Title", "Page"],
          [[f"Table {n}", C.TABLE_TITLES[n], ""] for n in sorted(C.TABLE_TITLES)],
          [1.0, 3.5, 1.0], number="", title="",
          )

    h1(doc, "LIST OF ABBREVIATIONS")
    table(doc,
          ["Abbreviation", "Full Form"],
          [[a, e] for a, e in C.ABBREVIATIONS],
          [1.6, 4.4], number="", title="")

    # ---------------------------------------------------------------- body --
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_numbering(body, "decimal", start=1)
    footer_page_number(body, show=True)

    # ------------------------------ Chapter 1 ------------------------------
    h1(doc, "CHAPTER 1: INTRODUCTION", page_break=False)
    h2(doc, "1.1 Introduction")
    for t in C.CH1_INTRO:
        para(doc, t)
    h2(doc, "1.2 Problem Statement")
    for t in C.CH1_PROBLEM_INTRO:
        para(doc, t)
    for t in C.CH1_PROBLEM_POINTS:
        bullet(doc, t)
    h2(doc, "1.3 Objectives")
    sub(doc, "General Objective:")
    bullet(doc, C.CH1_GENERAL_OBJECTIVE)
    sub(doc, "Specific Objectives:")
    for i, t in enumerate(C.CH1_SPECIFIC_OBJECTIVES, 1):
        numbered(doc, i, t)
    para(doc, "The third objective is measurable and verifiable: the modules named are "
              "either present or absent, and the 40 test cases described in Chapter 5 "
              "either pass or fail.")
    h2(doc, "1.4 Scope")
    for t in C.CH1_SCOPE_INTRO:
        para(doc, t)
    for t in C.CH1_SCOPE_POINTS:
        bullet(doc, t)
    h2(doc, "1.5 Significance and Limitation")
    sub(doc, "Significance:")
    for t in C.CH1_SIGNIFICANCE:
        bullet(doc, t)
    sub(doc, "Limitation:")
    for t in C.CH1_LIMITATION:
        bullet(doc, t)

    # ------------------------------ Chapter 2 ------------------------------
    h1(doc, "CHAPTER 2: LITERATURE REVIEW")
    h2(doc, "2.1 Research Based on Similar Projects")
    for t in C.CH2_RESEARCH:
        para(doc, t)
    table(doc,
          ["System", "Deployment", "Approach", "Multi-tenant", "Limitation"],
          [list(r) for r in C.TABLE1_ROWS],
          [1.0, 0.85, 1.55, 0.8, 1.8],
          number=1, title=C.TABLE_TITLES[1])
    h2(doc, "2.2 Related Theory")
    para(doc, C.CH2_THEORY_INTRO)
    for lead, text in C.CH2_THEORY_ITEMS:
        bullet(doc, text, lead=lead)

    # ------------------------------ Chapter 3 ------------------------------
    h1(doc, "CHAPTER 3: SYSTEM ANALYSIS")
    h2(doc, "3.1 Requirement Analysis")
    para(doc, C.CH3_REQ_INTRO)
    h3(doc, "3.1.1 Functional Requirements")
    para(doc, "These are the key functions the system must perform.")
    for title, items in C.CH3_FUNCTIONAL:
        sub(doc, title)
        for it in items:
            bullet(doc, it)
    h3(doc, "3.1.2 Non-Functional Requirements")
    para(doc, "These requirements define how well the system must perform.")
    for title, items in C.CH3_NONFUNCTIONAL:
        sub(doc, title)
        for it in items:
            bullet(doc, it)
    h2(doc, "3.2 Feasibility Study")
    para(doc, C.CH3_FEAS_INTRO)
    for heading, items in (("3.2.1 Technical Feasibility", C.CH3_TECHNICAL),
                           ("3.2.2 Operational Feasibility", C.CH3_OPERATIONAL),
                           ("3.2.3 Economic Feasibility", C.CH3_ECONOMIC),
                           ("3.2.4 Scheduling Feasibility", C.CH3_SCHEDULE)):
        h3(doc, heading)
        para(doc, items[0])
        for it in items[1:]:
            bullet(doc, it)

    # ------------------------------ Chapter 4 ------------------------------
    h1(doc, "CHAPTER 4: SYSTEM DESIGN AND ARCHITECTURE")
    h2(doc, "4.1 System Overview")
    for t in C.CH4_OVERVIEW:
        para(doc, t)
    h2(doc, "4.2 Selected SDLC Model")
    for t in C.CH4_SDLC:
        para(doc, t)
    figure(doc, 1)
    h2(doc, "4.3 System Requirements")
    h3(doc, "4.3.1 Hardware Requirements")
    table(doc, ["Component", "Specification", "Purpose"],
          [list(r) for r in C.TABLE2_ROWS], [1.4, 2.2, 2.4],
          number=2, title=C.TABLE_TITLES[2])
    h3(doc, "4.3.2 Software Requirements")
    table(doc, ["Layer", "Tool or Technology", "Purpose"],
          [list(r) for r in C.TABLE3_ROWS], [1.4, 2.2, 2.4],
          number=3, title=C.TABLE_TITLES[3])
    h2(doc, "4.4 Multi-Tenancy Model and Data Isolation")
    for t in C.CH4_TENANCY:
        para(doc, t)
    h2(doc, "4.5 Core Algorithms")
    para(doc, C.CH4_ALGO_INTRO)
    for heading, intro, steps in (C.ALGO_1, C.ALGO_2, C.ALGO_3):
        h3(doc, heading)
        para(doc, intro)
        for i, s in enumerate(steps, 1):
            numbered(doc, i, s)
    h2(doc, "4.6 Proposed System Architecture")
    h3(doc, "4.6.1 Overall System Architecture")
    for t in C.CH4_ARCH_OVERALL:
        para(doc, t)
    figure(doc, 2)
    h3(doc, "4.6.2 Component Description")
    para(doc, "The responsibilities of each component shown in Figure 2 are as follows.")
    for lead, text in C.CH4_COMPONENTS:
        bullet(doc, text, lead=lead)
    h3(doc, "4.6.3 Workflow Diagram")
    para(doc, C.CH4_WORKFLOW[0])
    figure(doc, 3)
    for t in C.CH4_WORKFLOW[1:]:
        para(doc, t)
    h2(doc, "4.7 System Design Diagrams")
    h3(doc, "4.7.1 Context Diagram")
    for t in C.CH4_CONTEXT:
        para(doc, t)
    figure(doc, 4)
    h3(doc, "4.7.2 Level 1 Data Flow Diagram")
    para(doc, C.CH4_DFD1[0])
    figure(doc, 5)
    for t in C.CH4_DFD1[1:]:
        para(doc, t)
    h3(doc, "4.7.3 Use Case Diagram")
    para(doc, C.CH4_USECASE[0])
    figure(doc, 6)
    for t in C.CH4_USECASE[1:]:
        para(doc, t)
    h3(doc, "4.7.4 ER Diagram")
    para(doc, C.CH4_ERD[0])
    figure(doc, 7)
    for t in C.CH4_ERD[1:]:
        para(doc, t)

    # ------------------------------ Chapter 5 ------------------------------
    h1(doc, "CHAPTER 5: EXPECTED OUTCOMES")
    h2(doc, "5.1 Expected Outcome")
    para(doc, C.CH5_INTRO)
    for t in C.CH5_OUTCOMES:
        bullet(doc, t)
    h2(doc, "5.2 Evaluation Criteria")
    para(doc, C.CH5_EVAL_INTRO)
    for heading, items in C.CH5_EVAL_SECTIONS:
        h3(doc, heading)
        for it in items:
            bullet(doc, it)
    h3(doc, "5.2.5 Evaluation Summary Table")
    table(doc, ["Criterion", "Target", "How it is measured"],
          [list(r) for r in C.TABLE4_ROWS], [2.1, 1.5, 2.4],
          number=4, title=C.TABLE_TITLES[4])
    h3(doc, "5.2.6 Acceptance Criteria")
    para(doc, C.CH5_ACCEPTANCE_INTRO)
    for t in C.CH5_ACCEPTANCE:
        bullet(doc, t)
    para(doc, C.CH5_ACCEPTANCE_TAIL)

    # ------------------------------ Chapter 6 ------------------------------
    h1(doc, "CHAPTER 6: TIME SCHEDULE")
    h2(doc, "6.1 Gantt Chart")
    para(doc, C.CH6_GANTT[0])
    figure(doc, 8)
    for t in C.CH6_GANTT[1:]:
        para(doc, t)

    # ------------------------------ References -----------------------------
    h1(doc, "REFERENCES")
    for i, ref in enumerate(C.REFERENCES, 1):
        p = doc.add_paragraph()
        _fmt(p, hanging=Inches(0.45), space_after=0,
             align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_run(p.add_run(f"[{i}]  {ref}"))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
