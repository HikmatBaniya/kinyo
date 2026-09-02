"""Builds the Kinyo Project-VI proposal as an APA 7th edition Word document.

Formatting implemented here (from the BIT Project-VI template):
  * Times New Roman 12 pt throughout; double spacing; 0 pt paragraph spacing
  * A4 portrait, 1 in margins, flush-left ragged-right text
  * 0.5 in first-line indent on every body paragraph
  * Page numbers top right; front matter in lower-case roman, body restarting
    at 1 in arabic
  * APA heading levels 1-3 by position and emphasis, all at 12 pt
  * Tables with horizontal rules only; table number bold above, italic title below
  * Figure number bold above, italic title below it, image beneath the title
  * Reference list alphabetical with a 0.5 in hanging indent

Run:  python build_docx.py
"""

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import content as C

# Page numbers for the List of Figures and List of Tables. Produced by
# index_pages.py from the exported PDF, so the first build of a fresh checkout
# leaves the column blank and the second fills it in.
_INDEX_FILE = Path(__file__).with_name("page_index.json")
if _INDEX_FILE.exists():
    PAGE_INDEX = json.loads(_INDEX_FILE.read_text(encoding="utf-8"))
else:
    PAGE_INDEX = {"figures": {}, "tables": {}}

HERE = Path(__file__).parent
FIG_DIR = HERE / "figures"
OUT = HERE.parent / "Kinyo_ProjectVI_Proposal.docx"

FONT = "Times New Roman"
SIZE = Pt(12)
INDENT = Inches(0.5)


# --------------------------------------------------------------- low level --
def _set_run(run, bold=False, italic=False, size=SIZE, name=FONT):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def _fmt(p, *, double=True, first_line=None, space_after=0, keep_next=False,
         align=WD_ALIGN_PARAGRAPH.LEFT, left_indent=None, hanging=None):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    if double:
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        pf.line_spacing = 1.0
    if first_line is not None:
        pf.first_line_indent = first_line
    if left_indent is not None:
        pf.left_indent = left_indent
    if hanging is not None:
        pf.left_indent = hanging
        pf.first_line_indent = -hanging
    pf.keep_with_next = keep_next
    p.alignment = align
    return p


def para(doc, text, *, indent=True, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
         italic=False, double=True, keep_next=False, left_indent=None,
         space_after=0, size=SIZE):
    p = doc.add_paragraph()
    _fmt(p, double=double, first_line=INDENT if indent else Inches(0),
         align=align, keep_next=keep_next, left_indent=left_indent,
         space_after=space_after)
    r = p.add_run(text)
    _set_run(r, bold=bold, italic=italic, size=size)
    return p


def blank(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        _fmt(p, double=False, first_line=Inches(0))
        _set_run(p.add_run(""))


def h1(doc, text, *, page_break=True):
    """APA level 1: centred, bold, title case, on a new page.

    Built-in Heading styles are used so that the table of contents field can
    find the headings; the APA appearance is then restored by direct
    formatting, which overrides the style.
    """
    p = doc.add_paragraph(style="Heading 1")
    _fmt(p, first_line=Inches(0), align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True)
    p.paragraph_format.page_break_before = page_break
    _set_run(p.add_run(text), bold=True)
    return p


def h2(doc, text):
    """APA level 2: flush left, bold, title case."""
    p = doc.add_paragraph(style="Heading 2")
    _fmt(p, first_line=Inches(0), keep_next=True)
    _set_run(p.add_run(text), bold=True)
    return p


def h3(doc, text):
    """APA level 3: flush left, bold italic, title case."""
    p = doc.add_paragraph(style="Heading 3")
    _fmt(p, first_line=Inches(0), keep_next=True)
    _set_run(p.add_run(text), bold=True, italic=True)
    return p


def bullet(doc, text, *, bold_lead=None):
    p = doc.add_paragraph()
    _fmt(p, hanging=Inches(0.5))
    if bold_lead:
        _set_run(p.add_run("•\t"))
        _set_run(p.add_run(bold_lead + ": "), bold=True)
        _set_run(p.add_run(text))
    else:
        _set_run(p.add_run("•\t" + text))
    return p


def numbered(doc, idx, text):
    p = doc.add_paragraph()
    _fmt(p, hanging=Inches(0.5))
    _set_run(p.add_run(f"{idx}.\t{text}"))
    return p


# ------------------------------------------------------------------ fields --
def _field(paragraph, instr):
    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r._element.append(fld)

    r = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    r._element.append(it)

    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    r._element.append(fld)

    r = paragraph.add_run("1")
    _set_run(r)

    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    r._element.append(fld)


def page_number_header(section, *, show=True):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    _fmt(p, double=False, first_line=Inches(0), align=WD_ALIGN_PARAGRAPH.RIGHT)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if show:
        _field(p, " PAGE ")


def set_page_numbering(section, fmt, start=None):
    sectPr = section._sectPr
    for existing in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(existing)
    el = OxmlElement("w:pgNumType")
    el.set(qn("w:fmt"), fmt)
    if start is not None:
        el.set(qn("w:start"), str(start))
    sectPr.append(el)


# ------------------------------------------------------------------ tables --
def _borders(el, spec):
    """spec: dict edge -> 'single' | 'none'."""
    borders = OxmlElement("w:tblBorders" if el.tag.endswith("tblPr") else "w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in spec:
            continue
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), spec[edge])
        if spec[edge] != "none":
            e.set(qn("w:sz"), "8")
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), "000000")
        borders.append(e)
    el.append(borders)


def apa_table(doc, headers, rows, widths, *, number, title, note=None,
              shade_cells=None, size=Pt(12)):
    """Table number bold above, italic title below it, horizontal rules only."""
    p = doc.add_paragraph()
    _fmt(p, first_line=Inches(0), keep_next=True)
    _set_run(p.add_run(f"Table {number}"), bold=True)

    p = doc.add_paragraph()
    _fmt(p, first_line=Inches(0), keep_next=True)
    _set_run(p.add_run(title), italic=True)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tblPr = table._tbl.tblPr
    _borders(tblPr, {"top": "single", "bottom": "single", "left": "none",
                     "right": "none", "insideH": "none", "insideV": "none"})
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    for col, width in zip(table.columns, widths):
        col.width = Inches(width)

    def fill(cells, values, bold=False):
        for cell, value, width in zip(cells, values, widths):
            cell.width = Inches(width)
            cp = cell.paragraphs[0]
            _fmt(cp, double=False, first_line=Inches(0), space_after=2)
            _set_run(cp.add_run(str(value)), bold=bold, size=size)

    fill(table.rows[0].cells, headers, bold=True)
    for cell in table.rows[0].cells:
        _borders(cell._tc.get_or_add_tcPr(), {"bottom": "single"})
    # repeat the header row when the table breaks across pages
    trPr = table.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)

    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        fill(cells, row)
        if shade_cells:
            for c_idx in shade_cells.get(r_idx, []):
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "BFBFBF")
                cells[c_idx]._tc.get_or_add_tcPr().append(shd)

    if note:
        p = doc.add_paragraph()
        _fmt(p, first_line=Inches(0))
        _set_run(p.add_run("Note. "), italic=True)
        _set_run(p.add_run(note))
    blank(doc)
    return table


def figure(doc, number, *, note=None):
    title, filename, width_in = C.FIGURES[number]
    p = doc.add_paragraph()
    _fmt(p, first_line=Inches(0), keep_next=True)
    _set_run(p.add_run(f"Figure {number}"), bold=True)

    p = doc.add_paragraph()
    _fmt(p, first_line=Inches(0), keep_next=True)
    _set_run(p.add_run(title), italic=True)

    p = doc.add_paragraph()
    _fmt(p, double=False, first_line=Inches(0), align=WD_ALIGN_PARAGRAPH.LEFT)
    p.add_run().add_picture(str(FIG_DIR / filename), width=Inches(width_in))

    if note:
        p = doc.add_paragraph()
        _fmt(p, first_line=Inches(0))
        _set_run(p.add_run("Note. "), italic=True)
        _set_run(p.add_run(note))
    blank(doc)


# ------------------------------------------------------------------- build --
def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = SIZE
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    rpr = style.element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)
    rpr.append(rfonts)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        hs = doc.styles[name]
        hs.font.name = FONT
        hs.font.size = SIZE
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.bold = True
        hs.paragraph_format.space_before = Pt(0)
        hs.paragraph_format.space_after = Pt(0)
        hs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        hrpr = hs.element.get_or_add_rPr()
        hrfonts = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            hrfonts.set(qn(attr), FONT)
        hrpr.append(hrfonts)

    sec = doc.sections[0]
    sec.page_width = Inches(8.27)
    sec.page_height = Inches(11.69)
    for side in ("left", "right", "top", "bottom"):
        setattr(sec, f"{side}_margin", Inches(1))
    page_number_header(sec, show=False)

    m = C.META

    # ------------------------------------------------------- title page --
    def centre(text, bold=False, italic=False, space=0):
        p = doc.add_paragraph()
        _fmt(p, double=False, first_line=Inches(0),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=space)
        _set_run(p.add_run(text), bold=bold, italic=italic)

    logo = FIG_DIR / "pu_logo.png"
    if logo.exists():
        p = doc.add_paragraph()
        _fmt(p, double=False, first_line=Inches(0),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
        p.add_run().add_picture(str(logo), width=Inches(1.39), height=Inches(1.22))
    else:
        blank(doc, 2)
    centre(m["university"], bold=True)
    centre(m["college"], bold=True)
    centre(m["college_address"], space=10)
    blank(doc)
    centre("A")
    centre("Project-VI Proposal Report")
    centre("on", space=6)
    centre(m["title"], bold=True, space=4)
    centre(m["semester_line"], space=14)
    blank(doc)
    centre("Submitted by", space=6)
    for name, roll in m["students"]:
        centre(f"{name} [{roll}]")
    blank(doc)
    centre("Under the Supervision of", space=6)
    centre(m["supervisor"], space=14)
    blank(doc)
    centre("Submitted to the Department of Science and Technology in Partial Fulfilment")
    centre("of the Requirements for the Degree of Bachelor of Information Technology",
           space=14)
    blank(doc)
    centre("Department of Science and Technology")
    centre(f'{m["city"]}, Nepal')
    centre(m["month_year"])

    # ------------------------------------------------------ front matter --
    fm = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_numbering(fm, "lowerRoman", start=2)
    page_number_header(fm, show=True)

    h1(doc, "Table of Contents", page_break=False)
    p = doc.add_paragraph()
    _fmt(p, first_line=Inches(0))
    _field(p, r' TOC \o "1-3" \h \z \u ')

    h1(doc, "List of Figures")
    apa_table(
        doc,
        ["Figure", "Title", "Page"],
        [[str(n), C.FIGURES[n][0], PAGE_INDEX["figures"].get(str(n), "")]
         for n in sorted(C.FIGURES)],
        [0.8, 4.6, 0.8],
        number="—", title="List of Figures",
        note="Figures are numbered consecutively in the order of first mention, as "
             "required by APA 7th edition.")

    h1(doc, "List of Tables")
    apa_table(
        doc,
        ["Table", "Title", "Page"],
        [[str(n), C.TABLE_TITLES[n], PAGE_INDEX["tables"].get(str(n), "")]
         for n in sorted(C.TABLE_TITLES)],
        [0.8, 4.6, 0.8],
        number="—", title="List of Tables",
        note="Tables are numbered consecutively in the order of first mention, as "
             "required by APA 7th edition.")

    h1(doc, "List of Abbreviations")
    apa_table(
        doc,
        ["Abbreviation", "Expansion"],
        [[a, e] for a, e in C.ABBREVIATIONS],
        [1.5, 4.7],
        number="—", title="List of Abbreviations")

    # -------------------------------------------------------------- body --
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_numbering(body, "decimal", start=1)
    page_number_header(body, show=True)

    # ---------------- Chapter 1 ----------------
    h1(doc, "Chapter 1: Introduction", page_break=False)
    h2(doc, "1.1 Background")
    for t in C.CH1_BACKGROUND:
        para(doc, t)
    h2(doc, "1.2 Problem Statement")
    for t in C.CH1_PROBLEM:
        para(doc, t)
    h2(doc, "1.3 Objectives")
    h3(doc, "1.3.1 General Objective")
    para(doc, C.CH1_GENERAL_OBJECTIVE)
    h3(doc, "1.3.2 Specific Objectives")
    para(doc, "The specific objectives of the project are the following.")
    for i, t in enumerate(C.CH1_SPECIFIC_OBJECTIVES, 1):
        numbered(doc, i, t)
    para(doc, "The third objective is measurable and verifiable: the modules named are "
              "either present or absent, and the 40 test cases described in Section 4.3 "
              "either pass or fail. No performance improvement is claimed, because no "
              "baseline has been measured.")
    h2(doc, "1.4 Project Scope")
    para(doc, C.CH1_SCOPE_INTRO)
    for lead, text in C.CH1_SCOPE_ITEMS:
        bullet(doc, text, bold_lead=lead)

    # ---------------- Chapter 2 ----------------
    h1(doc, "Chapter 2: Literature Overview")
    h2(doc, "2.1 Background and Theoretical Framework")
    for t in C.CH2_THEORY:
        para(doc, t)

    h2(doc, "2.2 Study of Related Systems and Related Work")
    para(doc, "Six existing systems relevant to the problem area were reviewed. Table 1 "
              "compares them by platform, technique, key features and limitations. Each "
              "entry is cited in the reference list.")
    apa_table(
        doc,
        ["Author / System", "Platform", "Technique or Approach", "Key Features",
         "Limitations"],
        [list(r) for r in C.TABLE1_ROWS],
        [1.0, 0.9, 1.35, 1.35, 1.6],
        number=1, title=C.TABLE_TITLES[1],
        note="All systems listed were reviewed through their official documentation, "
             "which is cited in the reference list.")
    for t in C.CH2_TABLE1_DISCUSSION:
        para(doc, t)

    h2(doc, "2.3 Contribution of the Proposed System")
    para(doc, C.CH2_CONTRIBUTION_INTRO)
    for t in C.CH2_CONTRIBUTION_ITEMS:
        bullet(doc, t)

    h2(doc, "2.4 Functional and Non-Functional Requirements")
    para(doc, "Table 2 states what the system must do and how well it must perform. "
              "Priority is recorded as High for requirements without which the system "
              "cannot meet its general objective, and Medium for those that improve the "
              "system but are not essential to a working release.")
    apa_table(
        doc,
        ["Requirement Type", "Description", "Priority"],
        [list(r) for r in C.TABLE2_ROWS],
        [1.8, 3.6, 0.8],
        number=2, title=C.TABLE_TITLES[2])

    h2(doc, "2.5 Feasibility Study")
    para(doc, "The feasibility study evaluates whether the proposed project can be "
              "completed with the available resources, time and skills. Table 3 "
              "summarises the analysis.")
    apa_table(
        doc,
        ["Feasibility Type", "Analysis", "Verdict"],
        [list(r) for r in C.TABLE3_ROWS],
        [1.35, 4.05, 0.8],
        number=3, title=C.TABLE_TITLES[3])
    para(doc, "Based on the above analysis, the project is determined to be feasible "
              "within the given academic timeframe, available tools and team capability.")

    # ---------------- Chapter 3 ----------------
    h1(doc, "Chapter 3: System Design and Methodology")
    h2(doc, "3.1 Software Development Life Cycle")
    for t in C.CH3_SDLC:
        para(doc, t)
    figure(doc, 1)

    h2(doc, "3.2 Project Timeline")
    para(doc, "Table 4 presents the proposed timeline. Requirement analysis, design, "
              "development, testing and documentation appear as separate activities, and "
              "testing is allocated its own period rather than being absorbed into "
              "development.")
    shade = {i: [1 + mi for mi in months] for i, (_a, months) in enumerate(C.TIMELINE)}
    apa_table(
        doc,
        ["Activity"] + C.MONTHS,
        [[a] + [""] * len(C.MONTHS) for a, _m in C.TIMELINE],
        [2.8] + [0.55] * len(C.MONTHS),
        number=4, title=C.TABLE_TITLES[4],
        shade_cells=shade,
        note="Shaded cells indicate the planned duration of each activity. Replace M1 to "
             "M6 with the actual month names of the project period.")

    h2(doc, "3.3 System Architecture")
    for t in C.CH3_ARCHITECTURE:
        para(doc, t)
    figure(doc, 2)

    h2(doc, "3.4 Algorithm")
    para(doc, C.CH3_ALGO_INTRO)
    for title, intro, steps in (
        (C.ALGO_1_TITLE, C.ALGO_1_INTRO, C.ALGO_1_STEPS),
        (C.ALGO_2_TITLE, C.ALGO_2_INTRO, C.ALGO_2_STEPS),
        (C.ALGO_3_TITLE, C.ALGO_3_INTRO, C.ALGO_3_STEPS),
    ):
        h3(doc, title)
        para(doc, intro)
        for i, s in enumerate(steps, 1):
            numbered(doc, i, s)

    h2(doc, "3.5 System Flowchart")
    para(doc, C.CH3_FLOWCHART[0])
    figure(doc, 3,
           note="Numbered circles are off-page connectors: control leaves a connector in "
                "one stage and resumes at the connector with the same number in another "
                "stage.")
    for t in C.CH3_FLOWCHART[1:]:
        para(doc, t)

    h2(doc, "3.6 Use Case Diagram")
    para(doc, C.CH3_USECASE[0])
    figure(doc, 4)
    for t in C.CH3_USECASE[1:]:
        para(doc, t)

    h2(doc, "3.7 Data Flow Diagram")
    h3(doc, "3.7.1 Level 0: Context Diagram")
    for t in C.CH3_DFD0:
        para(doc, t)
    figure(doc, 5)
    h3(doc, "3.7.2 Level 1: Detailed Data Flow Diagram")
    para(doc, C.CH3_DFD1[0])
    figure(doc, 6)
    for t in C.CH3_DFD1[1:]:
        para(doc, t)

    h2(doc, "3.8 Entity Relationship Diagram")
    para(doc, C.CH3_ERD[0])
    figure(doc, 7,
           note="Every tenant-owned entity carries tenant_id as a foreign key referencing "
                "TENANT. For legibility only the principal ownership relationships are "
                "drawn.")
    for t in C.CH3_ERD[1:]:
        para(doc, t)

    # ---------------- Chapter 4 ----------------
    h1(doc, "Chapter 4: Implementation Plan")
    para(doc, C.CH4_INTRO)
    h2(doc, "4.1 Hardware and Software Requirements")
    para(doc, "Table 5 lists the hardware and software required to develop, test and "
              "deploy the system.")
    apa_table(
        doc,
        ["Component", "Specification / Tool", "Purpose"],
        [list(r) for r in C.TABLE5_ROWS],
        [1.4, 2.2, 2.6],
        number=5, title=C.TABLE_TITLES[5])

    h2(doc, "4.2 Proposed Technology Stack")
    para(doc, "Table 6 lists the proposed technology stack. Each major choice is "
              "justified after the table.")
    apa_table(
        doc,
        ["Layer", "Proposed Choice"],
        [list(r) for r in C.TABLE6_ROWS],
        [2.3, 3.9],
        number=6, title=C.TABLE_TITLES[6])
    h3(doc, "Justification of the Technology Choices")
    for t in C.CH4_JUSTIFICATION:
        para(doc, t)

    h2(doc, "4.3 Proposed Testing and Development Approach")
    para(doc, C.CH4_STRUCTURE_INTRO)
    for lead, text in C.CH4_APPROACH_ITEMS:
        bullet(doc, text, bold_lead=lead)
    para(doc, "Git repositories:", indent=False)
    bullet(doc, m["repo_backend"], bold_lead="Back end")
    bullet(doc, m["repo_frontend"], bold_lead="Front end")
    para(doc, "The repository links above are maintained from the start of the project, "
              "as required by the proposal checklist.")

    # ---------------- Chapter 5 ----------------
    h1(doc, "Chapter 5: Expected Outcomes")
    para(doc, C.CH5_INTRO)
    h2(doc, "5.1 Expected System Deliverables")
    para(doc, "The project is expected to deliver the following.")
    for t in C.CH5_DELIVERABLES:
        bullet(doc, t)
    h2(doc, "5.2 Expected System Features and Benefits")
    for t in C.CH5_BENEFITS:
        para(doc, t)

    # ---------------- References ----------------
    h1(doc, "References")
    for ref in C.REFERENCES:
        p = doc.add_paragraph()
        _fmt(p, hanging=Inches(0.5))
        # segments wrapped in asterisks are titles of larger works, set in italics
        for i, seg in enumerate(ref.split("*")):
            if seg:
                _set_run(p.add_run(seg), italic=bool(i % 2))

    doc.save(OUT)
    print(f"wrote {OUT}")
    downloads = Path.home() / "Downloads" / "Kinyo_Project_Proposal.docx"
    try:
        doc.save(downloads)
        print(f"wrote {downloads}")
    except PermissionError:
        print(f"SKIPPED {downloads} — the file is open in Word; close it and "
              f"rerun to refresh the copy in Downloads")


if __name__ == "__main__":
    build()
