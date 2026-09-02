"""Builds the proposal by filling in the official BIT template rather than
recreating it.

Working from a copy of BIT_ProjectVI_Proposal_Template_General_APA.docx keeps
the university logo, the title page, the chapter and section headings and their
order exactly as the department issued them. Under each heading the template's
instruction text is removed and the project's own content inserted, formatted to
the rules stated in the template's own Tables A1-A4.

Two things the template does not itself provide and which this script adds:
  * page numbers in the top-right corner, roman in the front matter and
    restarting at 1 in arabic from Chapter 1
  * a 0.5 in first-line indent on body paragraphs

Run:  python fill_template.py
"""

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

import content as C
import content_prose as P

HERE = Path(__file__).parent
FIG_DIR = HERE / "figures"
TEMPLATE = Path.home() / "Downloads" / "BIT_ProjectVI_Proposal_Template_General_APA.docx"
OUT = HERE.parent / "Kinyo_ProjectVI_Proposal.docx"
DOWNLOADS = Path.home() / "Downloads" / "Kinyo_Project_Proposal.docx"

_INDEX = HERE / "page_index.json"
PAGE_INDEX = (json.loads(_INDEX.read_text(encoding="utf-8"))
              if _INDEX.exists() else {"figures": {}, "tables": {}})

FONT = "Times New Roman"
SIZE = Pt(12)
HEADING_SIZES = {"Heading 1": Pt(16), "Heading 2": Pt(14),
                 "Heading 3": Pt(14)}
INDENT = Inches(0)     # the report is set without first-line indents


# ---------------------------------------------------------------- helpers --
def set_run(run, *, bold=False, italic=False, size=SIZE):
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), FONT)


def fmt(p, *, double=True, first_line=None, align=WD_ALIGN_PARAGRAPH.LEFT,
        hanging=None, keep_next=False, space_after=0):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    if double:
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        pf.line_spacing = 1.0
    if hanging is not None:
        pf.left_indent = hanging
        pf.first_line_indent = -hanging
    elif first_line is not None:
        pf.first_line_indent = first_line
    pf.keep_with_next = keep_next
    p.alignment = align
    return p


class Cursor:
    """Inserts new block-level content after a moving anchor element."""

    def __init__(self, doc, anchor_el):
        self.doc = doc
        self.el = anchor_el

    def _attach(self, block_el):
        self.el.addnext(block_el)
        self.el = block_el

    def para(self, text, *, indent=True, italic=False, bold=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, double=True, keep_next=False,
             space_after=0, size=SIZE):
        p = self.doc.add_paragraph()
        fmt(p, double=double, first_line=INDENT if indent else Inches(0),
            align=align, keep_next=keep_next, space_after=space_after)
        set_run(p.add_run(text), bold=bold, italic=italic, size=size)
        self._attach(p._element)
        return p

    def bullet(self, text, *, lead=None):
        p = self.doc.add_paragraph()
        fmt(p, hanging=Inches(0.5))
        p.paragraph_format.first_line_indent = Inches(-0.25)
        set_run(p.add_run("•\t"))
        if lead:
            set_run(p.add_run(lead + ": "), bold=True)
        set_run(p.add_run(text))
        self._attach(p._element)
        return p

    def numbered(self, idx, text):
        p = self.doc.add_paragraph()
        fmt(p, hanging=Inches(0.5))
        p.paragraph_format.first_line_indent = Inches(-0.25)
        set_run(p.add_run(f"{idx}.\t{text}"))
        self._attach(p._element)
        return p

    def blank(self):
        p = self.doc.add_paragraph()
        fmt(p, double=False, first_line=Inches(0))
        set_run(p.add_run(""))
        self._attach(p._element)
        return p

    def groups(self, blocks):
        """A bold run-in heading followed by its bullets, as in the sample."""
        for heading, items in blocks:
            p = self.doc.add_paragraph()
            fmt(p, first_line=Inches(0), keep_next=True)
            set_run(p.add_run(heading), bold=True)
            self._attach(p._element)
            for item in items:
                self.bullet(item)

    # --- figures and tables, per the template's Table A3 -------------------
    def figure(self, number, *, note=None):
        title, filename, width_in = C.FIGURES[number]
        if filename is None:            # placeholder for a figure supplied later
            p = self.doc.add_paragraph()
            fmt(p, double=False, first_line=Inches(0),
                align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True)
            set_run(p.add_run("[ Insert Gantt chart here ]"))
            self._attach(p._element)
            p = self.doc.add_paragraph()
            fmt(p, double=False, first_line=Inches(0),
                align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
            set_run(p.add_run(f"Figure {number} "), bold=True)
            set_run(p.add_run(title))
            self._attach(p._element)
            self.blank()
            return
        p = self.doc.add_paragraph()
        fmt(p, double=False, first_line=Inches(0),
            align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True)
        p.add_run().add_picture(str(FIG_DIR / filename), width=Inches(width_in))
        self._attach(p._element)

        # caption beneath the figure, number and title on one line
        p = self.doc.add_paragraph()
        fmt(p, double=False, first_line=Inches(0),
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        set_run(p.add_run(f"Figure {number} "), bold=True)
        set_run(p.add_run(title))
        self._attach(p._element)

        if note:
            p = self.doc.add_paragraph()
            fmt(p, first_line=Inches(0))
            set_run(p.add_run("Note. "))
            set_run(p.add_run(note))
            self._attach(p._element)
        self.blank()

    def table(self, headers, rows, widths, *, number, title, note=None,
              shade=None, size=Pt(12), caption=True):
        if caption:
            p = self.doc.add_paragraph()
            fmt(p, first_line=Inches(0), keep_next=True)
            set_run(p.add_run(f"Table {number}"), bold=True)
            self._attach(p._element)

            p = self.doc.add_paragraph()
            fmt(p, first_line=Inches(0), keep_next=True)
            set_run(p.add_run(title))
            self._attach(p._element)

        t = self.doc.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.autofit = False
        tblPr = t._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge, val in (("top", "single"), ("left", "none"), ("bottom", "single"),
                          ("right", "none"), ("insideH", "none"), ("insideV", "none")):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), val)
            if val != "none":
                e.set(qn("w:sz"), "8")
                e.set(qn("w:space"), "0")
                e.set(qn("w:color"), "000000")
            borders.append(e)
        tblPr.append(borders)
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)
        tw = OxmlElement("w:tblW")
        tw.set(qn("w:w"), str(int(sum(widths) * 1440)))
        tw.set(qn("w:type"), "dxa")
        tblPr.append(tw)
        for col, w in zip(t.columns, widths):
            col.width = Inches(w)

        def fill_row(cells, values, bold=False):
            for cell, value, w in zip(cells, values, widths):
                cell.width = Inches(w)
                cp = cell.paragraphs[0]
                fmt(cp, double=False, first_line=Inches(0), space_after=2)
                set_run(cp.add_run(str(value)), bold=bold, size=size)

        fill_row(t.rows[0].cells, headers, bold=True)
        for cell in t.rows[0].cells:
            tcb = OxmlElement("w:tcBorders")
            e = OxmlElement("w:bottom")
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "8")
            e.set(qn("w:color"), "000000")
            tcb.append(e)
            cell._tc.get_or_add_tcPr().append(tcb)
        trPr = t.rows[0]._tr.get_or_add_trPr()
        th = OxmlElement("w:tblHeader")
        th.set(qn("w:val"), "true")
        trPr.append(th)

        for r_i, row in enumerate(rows):
            cells = t.add_row().cells
            fill_row(cells, row)
            if shade:
                for c_i in shade.get(r_i, []):
                    sh = OxmlElement("w:shd")
                    sh.set(qn("w:val"), "clear")
                    sh.set(qn("w:fill"), "BFBFBF")
                    cells[c_i]._tc.get_or_add_tcPr().append(sh)

        self._attach(t._tbl)
        if note:
            p = self.doc.add_paragraph()
            fmt(p, first_line=Inches(0))
            set_run(p.add_run("Note. "))
            set_run(p.add_run(note))
            self._attach(p._element)
        self.blank()
        return t


def field(paragraph, instr, placeholder="1"):
    for text, kind in ((None, "begin"), (instr, None), (None, "separate"),
                       (placeholder, "text"), (None, "end")):
        r = paragraph.add_run()
        if kind in ("begin", "separate", "end"):
            f = OxmlElement("w:fldChar")
            f.set(qn("w:fldCharType"), kind)
            r._element.append(f)
        elif kind is None:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = text
            r._element.append(it)
        else:
            r.text = text
            set_run(r)


# ------------------------------------------------------- template surgery --
def body_children(doc):
    return list(doc.element.body.iterchildren())


def all_paragraphs(doc):
    return [Paragraph(el, doc) for el in doc.element.body.iterchildren()
            if el.tag == qn("w:p")]


def find_para(doc, prefix, style=None):
    for p in all_paragraphs(doc):
        if p.text.strip().startswith(prefix) and (style is None or p.style.name == style):
            return p
    raise LookupError(f"paragraph not found: {prefix!r}")


def is_heading(el, doc):
    if el.tag != qn("w:p"):
        return False
    return Paragraph(el, doc).style.name.startswith("Heading")


def clear_until_next_heading(doc, heading_el):
    """Delete every block after `heading_el` up to the next heading.

    Stops at the body-level w:sectPr, which closes the document and must never
    be removed: deleting it destroys the final section and with it the body page
    numbering.
    """
    el = heading_el.getnext()
    while el is not None and not is_heading(el, doc):
        if el.tag == qn("w:sectPr"):
            break
        nxt = el.getnext()
        el.getparent().remove(el)
        el = nxt


def delete_range(start_el, end_el):
    """Delete start_el..end_el inclusive."""
    el = start_el
    while el is not None:
        nxt = el.getnext()
        stop = el is end_el
        el.getparent().remove(el)
        if stop:
            return
        el = nxt


def set_text(paragraph, text, *, bold=None, italic=False, size=SIZE):
    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)
    run = paragraph.add_run(text)
    set_run(run, bold=bool(bold), italic=italic, size=size)
    return paragraph


def add_page_numbering(doc):
    """Top-right page numbers, roman for front matter, arabic from Chapter 1."""
    def header_number(section, show=True):
        # the number sits centred in the footer; any header is left empty
        section.header.is_linked_to_previous = False
        hp = section.header.paragraphs[0]
        fmt(hp, double=False, first_line=Inches(0))
        for r in list(hp.runs):
            r._element.getparent().remove(r._element)

        section.footer.is_linked_to_previous = False
        p = section.footer.paragraphs[0]
        fmt(p, double=False, first_line=Inches(0), align=WD_ALIGN_PARAGRAPH.CENTER)
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        if show:
            field(p, " PAGE ")

    def numbering(section, kind, start):
        sectPr = section._sectPr
        for old in sectPr.findall(qn("w:pgNumType")):
            sectPr.remove(old)
        el = OxmlElement("w:pgNumType")
        el.set(qn("w:fmt"), kind)
        el.set(qn("w:start"), str(start))
        # OOXML fixes the order of sectPr children; pgNumType must precede
        # w:cols, and Word silently drops it if it is simply appended.
        anchor = sectPr.find(qn("w:cols"))
        if anchor is not None:
            anchor.addprevious(el)
        else:
            sectPr.append(el)

    # section break before the Table of Contents and before Chapter 1
    for anchor_prefix, kind, start in (("Abstract", "lowerRoman", 2),
                                       ("Chapter 1:", "decimal", 1)):
        heading = find_para(doc, anchor_prefix)
        sectPr = OxmlElement("w:sectPr")
        e = OxmlElement("w:type")
        e.set(qn("w:val"), "nextPage")
        sectPr.append(e)

        # Carry the break on the preceding paragraph where there is one. Adding a
        # fresh empty paragraph for it can spill onto a page of its own and leave
        # a numbered blank page before the next section.
        prev = heading._element.getprevious()
        if prev is not None and prev.tag == qn("w:p"):
            carrier = Paragraph(prev, doc)
        else:
            carrier = doc.add_paragraph()
            fmt(carrier, double=False, first_line=Inches(0))
            heading._element.addprevious(carrier._element)
        carrier._element.get_or_add_pPr().append(sectPr)
        heading.paragraph_format.page_break_before = False

    base = doc.sections[0]
    for sec in doc.sections:
        sec.page_width, sec.page_height = base.page_width, base.page_height
        for side in ("left", "right", "top", "bottom"):
            setattr(sec, f"{side}_margin", Inches(1))

    # The template ships with sections of its own, so the two inserted breaks do
    # not land at fixed indices. The last section is always the body and the one
    # before it the front matter, counting from the end.
    secs = doc.sections
    header_number(secs[0], show=False)           # title page: no number printed
    numbering(secs[-2], "lowerRoman", 2)         # front matter
    header_number(secs[-2], show=True)
    numbering(secs[-1], "decimal", 1)            # body, restarting at 1
    header_number(secs[-1], show=True)
    return secs


def normalise_type(doc):
    """Times New Roman everywhere, upright: no italics anywhere in the report.

    The template ships its level-3 headings in bold italic and leaves the
    document default font as Calibri, which Word then uses for the generated
    table of contents. Both are overridden here.
    """
    rpr = doc.styles.element.find(qn("w:docDefaults"))
    if rpr is not None:
        for rf in rpr.iter(qn("w:rFonts")):
            for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rf.set(qn(a), FONT)

    for style in doc.styles:
        try:
            if style.font is None:
                continue
        except (AttributeError, NotImplementedError):
            continue
        style.font.italic = False
        if style.name and (style.name.lower().startswith(("toc", "heading"))
                           or style.name in ("normal", "Normal")):
            style.font.name = FONT

    def strip(paragraphs):
        for para in paragraphs:
            for r in para.runs:
                r.italic = False
                rpr = r._element.get_or_add_rPr()
                rf = rpr.find(qn("w:rFonts"))
                if rf is None:
                    rf = OxmlElement("w:rFonts")
                    rpr.append(rf)
                for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                    rf.set(qn(a), FONT)

    for name, size in HEADING_SIZES.items():
        doc.styles[name].font.size = size

    def resize_headings(paragraphs):
        for para in paragraphs:
            size = HEADING_SIZES.get(para.style.name)
            if size is None:
                continue
            for r in para.runs:
                r.font.size = size
                r.bold = True

    resize_headings(doc.paragraphs)

    strip(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                strip(cell.paragraphs)


def h3_after(doc, cursor, text):
    """Insert an APA level-3 heading (flush left, bold italic) after the cursor."""
    p = doc.add_paragraph(style="Heading 3")
    fmt(p, first_line=Inches(0), keep_next=True)
    set_run(p.add_run(text), bold=True)
    cursor._attach(p._element)
    return p


def section(doc, heading_prefix):
    """Clear a section's placeholder text and return a cursor after its heading."""
    h = find_para(doc, heading_prefix)
    clear_until_next_heading(doc, h._element)
    return Cursor(doc, h._element)


def build():
    doc = Document(str(TEMPLATE))
    m = C.META

    # ---- title page: fill the placeholders, keep the logo and the layout ----
    set_text(find_para(doc, "[NAME OF COLLEGE]"), m["college"], bold=True)
    set_text(find_para(doc, "[Address, City]"), m["college_address"])
    title_para = find_para(doc, "[Place Your Project Title Here]")
    set_text(title_para, m["title"], bold=True)
    semester = doc.add_paragraph()
    fmt(semester, double=False, first_line=Inches(0),
        align=WD_ALIGN_PARAGRAPH.CENTER)
    semester.paragraph_format.line_spacing = 1.5
    set_run(semester.add_run(m["semester_line"]))
    title_para._element.addnext(semester._element)
    students = [p for p in all_paragraphs(doc)
                if p.text.strip().startswith("Student Name")]
    for p, (name, roll) in zip(students, m["students"]):
        set_text(p, f"{name} [{roll}]")
    for leftover in students[len(m["students"]):]:
        leftover._element.getparent().remove(leftover._element)
    # supervisor block dropped from the title page at the authors' request
    for prefix in ("[Supervisor Name]", "Under the Supervision of"):
        para_el = find_para(doc, prefix)._element
        para_el.getparent().remove(para_el)
    set_text(find_para(doc, "[City], Nepal"), f'{m["city"]}, Nepal')
    set_text(find_para(doc, "[Month], [Year]"), m["month_year"])

    # The title page is double spaced in the template and overflows by one line
    # once the logo is included, so the front page alone is set to 1.5.
    toc_heading = find_para(doc, "Table of Contents")._element
    for el in doc.element.body.iterchildren():
        if el is toc_heading:
            break
        if el.tag == qn("w:p"):
            Paragraph(el, doc).paragraph_format.line_spacing = 1.5

    # ---- drop the two reference pages the template itself says to remove ----
    start = find_para(doc, "Proposal Report Formatting Guidelines")
    end = find_para(doc, "Remove this page from the final bound submission")
    delete_range(start._element, end._element)

    # ---- front matter ----
    toc = find_para(doc, "Table of Contents")
    abstract = doc.add_paragraph(style="Heading 1")
    fmt(abstract, first_line=Inches(0), align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True)
    set_run(abstract.add_run("Abstract"), bold=True)
    toc._element.addprevious(abstract._element)
    acur = Cursor(doc, abstract._element)
    for t in C.ABSTRACT:
        acur.para(t)
    toc.paragraph_format.page_break_before = True

    cur = section(doc, "Table of Contents")
    p = doc.add_paragraph()
    fmt(p, first_line=Inches(0))
    field(p, r' TOC \o "1-3" \h \z \u ', "Update this field in Word.")
    cur._attach(p._element)

    cur = section(doc, "List of Figures")
    cur.table(["Figure", "Title", "Page"],
              [[str(n), C.FIGURES[n][0], PAGE_INDEX["figures"].get(str(n), "")]
               for n in sorted(C.FIGURES)],
              [0.8, 4.6, 0.8], number=None, title=None, caption=False)

    if C.TABLE_TITLES:
        cur = section(doc, "List of Tables")
        cur.table(["Table", "Title", "Page"],
                  [[str(n), C.TABLE_TITLES[n], PAGE_INDEX["tables"].get(str(n), "")]
                   for n in sorted(C.TABLE_TITLES)],
                  [0.8, 4.6, 0.8], number=None, title=None, caption=False)
    else:
        # nothing is numbered as a table, so the section is dropped entirely
        heading = find_para(doc, "List of Tables")
        clear_until_next_heading(doc, heading._element)
        heading._element.getparent().remove(heading._element)

    cur = section(doc, "List of Abbreviations")
    cur.table(["Abbreviation", "Expansion"], [[a, e] for a, e in C.ABBREVIATIONS],
              [1.5, 4.7], number=None, title=None, caption=False)

    # ---- Chapter 1 ----
    cur = section(doc, "1.1 Background")
    for t in C.CH1_BACKGROUND:
        cur.para(t)
    cur = section(doc, "1.2 Problem Statement")
    for t in C.CH1_PROBLEM:
        cur.para(t)
    section(doc, "1.3 Objectives")
    cur = section(doc, "1.3.1 General Objective")
    cur.para(C.CH1_GENERAL_OBJECTIVE)
    cur = section(doc, "1.3.2 Specific Objectives")
    cur.para("The specific objectives of the project are the following.")
    for i, t in enumerate(C.CH1_SPECIFIC_OBJECTIVES, 1):
        cur.numbered(i, t)
    cur.para("The third objective is measurable and verifiable: the modules named are "
             "either present or absent, and the 40 test cases described in Section 4.3 "
             "either pass or fail. No performance improvement is claimed, because no "
             "baseline has been measured.")
    cur = section(doc, "1.4 Project Scope")
    cur.para(C.CH1_SCOPE_INTRO)
    for lead, text in C.CH1_SCOPE_ITEMS:
        cur.bullet(text, lead=lead)

    # ---- Chapter 2 ----
    cur = section(doc, "2.1 Background and Theoretical Framework")
    for t in C.CH2_THEORY:
        cur.para(t)
    cur = section(doc, "2.2 Study of Related Systems")
    cur.para("Six existing systems relevant to the problem area were reviewed. Each "
             "is summarised below by platform, approach, strengths and limitations, "
             "and each is cited in the reference list.")
    cur.groups(P.RELATED_SYSTEMS)
    for t in C.CH2_TABLE1_DISCUSSION:
        cur.para(t)
    cur = section(doc, "2.3 Contribution of the Proposed System")
    cur.para(C.CH2_CONTRIBUTION_INTRO)
    for t in C.CH2_CONTRIBUTION_ITEMS:
        cur.bullet(t)
    cur = section(doc, "2.4 Functional and Non-Functional Requirements")
    cur.para("This section states what the system must do and how well it must "
             "perform. The functional requirements describe the operations the "
             "system carries out; the non-functional requirements describe the "
             "qualities it must exhibit while doing so.")
    cur.para("Functional Requirements", indent=False, bold=True)
    cur.groups(P.FUNCTIONAL_REQUIREMENTS)
    cur.para("Non-Functional Requirements", indent=False, bold=True)
    cur.groups(P.NONFUNCTIONAL_REQUIREMENTS)
    cur = section(doc, "2.5 Feasibility Study")
    cur.para("The feasibility study evaluates whether the proposed platform can be "
             "developed and deployed within the available technical, operational, "
             "economic, schedule and legal constraints.")
    cur.groups(P.FEASIBILITY)
    cur.para("Based on the above analysis, the project is determined to be feasible "
             "within the given academic timeframe, available tools and team capability.")

    # ---- Chapter 3 ----
    cur = section(doc, "3.1 Software Development Life Cycle")
    for t in C.CH3_SDLC:
        cur.para(t)
    cur.figure(1)
    cur = section(doc, "3.2 Project Timeline")
    cur.para("The project runs for nineteen weeks, from Ashadh to Kartik. "
             "Requirement analysis, design, development, integration, testing, "
             "deployment and documentation are planned as separate activities, "
             "and testing is allocated its own period rather than being absorbed "
             "into development. The schedule is shown in Figure 2.")
    cur.groups([("Planned Activities", [
        "Requirement analysis and system design, five weeks in total, so that "
        "the data model and the tenancy strategy are settled before any "
        "application code is written.",
        "Database design and migration setup, two weeks, followed by the tenancy "
        "and authentication core over four weeks. This is the highest-risk part "
        "of the design and is built first so that every later increment "
        "exercises it.",
        "Catalogue and inventory, three weeks, then the storefront and "
        "host-based routing, three weeks.",
        "Cart and orders, four weeks, then the seller dashboard, administration "
        "console and reporting, two weeks.",
        "Integration of the modules, two weeks, overlapping the end of "
        "development.",
        "Testing, three weeks, covering unit, integration and user acceptance "
        "testing, followed by one week for deployment and domain configuration.",
        "Documentation and report writing run across all nineteen weeks, so that "
        "the final report does not depend on a single concentrated effort at the "
        "end.",
    ])])
    cur.figure(2)
    cur = section(doc, "3.3 System Architecture")
    for t in C.CH3_ARCHITECTURE:
        cur.para(t)
    cur.figure(3)
    cur = section(doc, "3.4 Algorithm")
    cur.para(C.CH3_ALGO_INTRO)
    for title, intro, steps in ((C.ALGO_1_TITLE, C.ALGO_1_INTRO, C.ALGO_1_STEPS),
                                (C.ALGO_2_TITLE, C.ALGO_2_INTRO, C.ALGO_2_STEPS),
                                (C.ALGO_3_TITLE, C.ALGO_3_INTRO, C.ALGO_3_STEPS)):
        h3_after(doc, cur, title)
        cur.para(intro)
        for i, s in enumerate(steps, 1):
            cur.numbered(i, s)
    cur = section(doc, "3.5 System Flowchart")
    cur.para(C.CH3_FLOWCHART[0])
    cur.figure(4)
    for t in C.CH3_FLOWCHART[1:]:
        cur.para(t)
    cur = section(doc, "3.6 Use Case Diagram")
    cur.para(C.CH3_USECASE[0])
    cur.figure(5)
    for t in C.CH3_USECASE[1:]:
        cur.para(t)
    section(doc, "3.7 Data Flow Diagram")
    cur = section(doc, "3.7.1 Level 0")
    for t in C.CH3_DFD0:
        cur.para(t)
    cur.figure(6)
    cur = section(doc, "3.7.2 Level 1")
    cur.para(C.CH3_DFD1[0])
    cur.figure(7)
    for t in C.CH3_DFD1[1:]:
        cur.para(t)
    cur = section(doc, "3.8 Entity Relationship Diagram")
    cur.para(C.CH3_ERD[0])
    cur.figure(8)
    for t in C.CH3_ERD[1:]:
        cur.para(t)

    # ---- Chapter 4 ----
    cur = section(doc, "Chapter 4: Implementation Plan")
    cur.para(C.CH4_INTRO)
    cur = section(doc, "4.1 Hardware and Software Requirements")
    cur.para("This section states the hardware and software required to develop, "
             "test and deploy the system.")
    cur.para("Hardware Requirements", indent=False, bold=True)
    cur.groups(P.HARDWARE_REQUIREMENTS)
    cur.para("Software Requirements", indent=False, bold=True)
    cur.groups(P.SOFTWARE_REQUIREMENTS)
    cur = section(doc, "4.2 Proposed Technology Stack")
    cur.para("The proposed technology stack is set out below, tier by tier, with "
             "the reason for each major choice stated alongside it.")
    cur.groups(P.TECHNOLOGY_STACK)
    cur = section(doc, "4.3 Proposed Testing and Development Approach")
    cur.para(C.CH4_STRUCTURE_INTRO)
    for lead, text in C.CH4_APPROACH_ITEMS:
        cur.bullet(text, lead=lead)
    cur.para("Git repositories:", indent=False)
    cur.bullet(m["repo_backend"], lead="Back end")
    cur.bullet(m["repo_frontend"], lead="Front end")
    cur.para("The repository links above are maintained from the start of the project, "
             "as required by the proposal checklist.")

    # ---- Chapter 5 ----
    cur = section(doc, "Chapter 5: Expected Outcomes")
    cur.para(C.CH5_INTRO)
    cur = section(doc, "5.1 Expected System Deliverables")
    cur.para("The project is expected to deliver the following.")
    for t in C.CH5_DELIVERABLES:
        cur.bullet(t)
    cur = section(doc, "5.2 Expected System Features and Benefits")
    for t in C.CH5_BENEFITS:
        cur.para(t)

    # ---- References ----
    cur = section(doc, "References")
    for ref in C.REFERENCES:
        p = doc.add_paragraph()
        fmt(p, hanging=Inches(0.5))
        set_run(p.add_run(ref.replace("*", "")))
        cur._attach(p._element)

    normalise_type(doc)

    # ---- page numbering, which the template does not itself supply ----
    add_page_numbering(doc)

    written = 0
    for target in (OUT, DOWNLOADS):
        try:
            doc.save(target)
            print(f"wrote {target}")
            written += 1
        except PermissionError:
            print(f"SKIPPED {target} - open in Word; close it and rerun")
    if not written:
        raise SystemExit("both output files are locked; close them in Word")


if __name__ == "__main__":
    build()
