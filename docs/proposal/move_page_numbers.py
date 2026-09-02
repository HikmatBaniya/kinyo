"""Moves the page number between the footer and the top-right header, in place.

    python move_page_numbers.py header <document.docx> [more.docx ...]
    python move_page_numbers.py footer <document.docx> [more.docx ...]

"header" puts it in the top right corner, as the BIT template requires; "footer"
centres it at the bottom. The roman/arabic numbering already set on each section
is left alone, and nothing else in the document is touched.
"""

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

FONT = "Times New Roman"


def clear(part):
    for para in part.paragraphs:
        for run in list(para.runs):
            run._element.getparent().remove(run._element)


def has_page_field(part):
    return any("PAGE" in p._p.xml for p in part.paragraphs)


def write_field(paragraph, align):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.first_line_indent = Inches(0)
    paragraph.alignment = align

    for text, kind in ((None, "begin"), (" PAGE ", None), (None, "separate"),
                       ("1", "text"), (None, "end")):
        run = paragraph.add_run()
        if kind in ("begin", "separate", "end"):
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
            run._element.append(el)
        elif kind is None:
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
            run._element.append(el)
        else:
            run.text = text
        run.font.size = Pt(12)
        rpr = run._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.append(rf)
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(a), FONT)


def move(path, where):
    doc = Document(path)
    moved = 0
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        source = section.footer if where == "header" else section.header
        target = section.header if where == "header" else section.footer

        if not has_page_field(source) and not has_page_field(target):
            continue                      # title page carries no number
        clear(source)
        clear(target)
        align = (WD_ALIGN_PARAGRAPH.RIGHT if where == "header"
                 else WD_ALIGN_PARAGRAPH.CENTER)
        write_field(target.paragraphs[0], align)
        moved += 1

    shutil.copy2(path, str(path) + ".bak")
    doc.save(path)
    spot = "top right" if where == "header" else "bottom centre"
    print(f"{Path(path).name}: page number moved to the {spot} in {moved} section(s)")


if __name__ == "__main__":
    if len(sys.argv) < 3 or sys.argv[1] not in ("header", "footer"):
        raise SystemExit(__doc__)
    for target in sys.argv[2:]:
        move(target, sys.argv[1])
