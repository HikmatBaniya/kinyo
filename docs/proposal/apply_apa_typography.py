"""Brings an already-built document back to the template's typography rules,
without touching its content or structure.

    python apply_apa_typography.py <document.docx> [more.docx ...]

Applies, per Tables A1-A3 of the BIT template:
  * every heading at 12 pt; level 3 bold italic
  * a 0.5 in first-line indent on body paragraphs
  * figure number bold above the image, italic title beneath the number
  * italics restored on the titles of larger works in the reference list

Bullets, numbered lists, reference hanging indents, table cells, the title page
and the front matter are left alone. The document is backed up first.
"""

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import content as C

FONT = "Times New Roman"
INDENT = Inches(0.5)


def style_run(run, *, bold=False, italic=False, size=Pt(12)):
    run.font.size = size
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), FONT)
    return run


def has_drawing(paragraph):
    return "graphicData" in paragraph._p.xml


def norm(text):
    return re.sub(r"\s+", " ", text).strip()


def fix_headings(doc):
    """Every heading at 12 pt; level 3 in bold italic, per Table A2."""
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        st = doc.styles[name]
        st.font.size = Pt(12)
        st.font.italic = (name == "Heading 3")
    count = 0
    for p in doc.paragraphs:
        if not p.style.name.startswith("Heading") or not p.text.strip():
            continue
        italic = p.style.name == "Heading 3"
        for r in p.runs:
            style_run(r, bold=True, italic=italic)
        count += 1
    return f"{count} headings set to 12 pt"


def fix_indents(doc):
    """0.5 in first line on body prose only."""
    count = 0
    for p in doc.paragraphs:
        pf = p.paragraph_format
        if p.style.name.startswith(("Heading", "toc", "TOC")):
            continue
        if not p.text.strip() or has_drawing(p):
            continue
        if p.alignment in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT):
            continue
        if pf.left_indent and pf.left_indent.inches > 0.01:
            continue                       # bullets, numbered items, references
        if pf.first_line_indent and pf.first_line_indent.inches != 0:
            continue
        pf.first_line_indent = INDENT
        count += 1
    return f"{count} body paragraphs indented"


def fix_figure_captions(doc):
    """Number bold above the image, italic title below it, image beneath."""
    wanted = {f"Figure {n} {title}": (n, title)
              for n, (title, _f, _w) in C.FIGURES.items()}
    moved = 0
    for p in list(doc.paragraphs):
        hit = wanted.get(norm(p.text))
        if not hit:
            continue
        number, title = hit
        prev = p._p.getprevious()
        if prev is None:
            continue
        anchor = prev                      # the paragraph holding the picture

        num_p = doc.add_paragraph()
        pf = num_p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0)
        pf.first_line_indent = Inches(0); pf.keep_with_next = True
        num_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style_run(num_p.add_run(f"Figure {number}"), bold=True)

        title_p = doc.add_paragraph()
        pf = title_p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0)
        pf.first_line_indent = Inches(0); pf.keep_with_next = True
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style_run(title_p.add_run(title), italic=True)

        anchor.addprevious(num_p._element)
        anchor.addprevious(title_p._element)
        p._p.getparent().remove(p._p)      # drop the caption from below
        moved += 1
    return f"{moved} figure captions moved above the image"


def fix_reference_italics(doc):
    """Restore italics on the titles of larger works."""
    marked = {norm(r.replace("*", "")): r for r in C.REFERENCES}
    count = 0
    for p in doc.paragraphs:
        source = marked.get(norm(p.text))
        if not source or "*" not in source:
            continue
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        for i, seg in enumerate(source.split("*")):
            if seg:
                style_run(p.add_run(seg), italic=bool(i % 2))
        count += 1
    return f"{count} references italicised"


def main(path):
    doc = Document(path)
    notes = [fix_headings(doc), fix_indents(doc),
             fix_figure_captions(doc), fix_reference_italics(doc)]
    shutil.copy2(path, str(path) + ".bak")
    doc.save(path)
    print(f"{Path(path).name}:")
    for n in notes:
        print(f"  - {n}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        raise SystemExit(__doc__)
    for target in sys.argv[1:]:
        main(target)
