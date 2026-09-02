"""Applies small edits to an already-built document without regenerating it.

Used when the .docx has been edited by hand: only the paragraphs named here are
touched and the rest of the file is left as it is.

    python patch_existing.py <document.docx> [more.docx ...]

Edits applied:
  * adds the semester line beneath the project title, if it is not there
  * removes the explanatory "Note." paragraphs beneath figures
"""

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

import content as C

FONT = "Times New Roman"


def style_run(run):
    run.font.size = Pt(12)
    run.bold = False
    run.italic = False
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), FONT)


def patch(path):
    doc = Document(path)
    changes = []

    # 1. semester line under the project title
    line = C.META["semester_line"]
    if not any(p.text.strip() == line for p in doc.paragraphs):
        title = next(p for p in doc.paragraphs
                     if p.text.strip() == C.META["title"])
        new = doc.add_paragraph()
        pf = new.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.5
        new.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(new.add_run(line))
        title._element.addnext(new._element)
        changes.append(f"added {line!r} under the title")

    # 2. drop the explanatory notes under figures
    removed = 0
    for p in list(doc.paragraphs):
        if p.text.strip().startswith("Note."):
            p._element.getparent().remove(p._element)
            removed += 1
    if removed:
        changes.append(f"removed {removed} Note. paragraph(s)")

    if not changes:
        print(f"{Path(path).name}: nothing to do")
        return

    shutil.copy2(path, str(path) + ".bak")
    doc.save(path)
    print(f"{Path(path).name}: " + "; ".join(changes))


if __name__ == "__main__":
    if len(sys.argv) < 2:
        raise SystemExit(__doc__)
    for target in sys.argv[1:]:
        patch(target)
